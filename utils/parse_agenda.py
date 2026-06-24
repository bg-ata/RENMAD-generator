# -*- coding: utf-8 -*-
"""
Parse a RENMAD/ATA event **agenda** (Word .docx — the primary input) into a
canonical agenda dict that drives title-slide generation.

The Word agenda is a single 2-column table:  C0 = time range, C1 = content block.
A content block is multiple lines:

    <session title>
    <Type> | Language: <English|Italian>      (optional — Discussion Panel / Presentation)
    <description paragraph>                     (ignored — never shown on a title slide)
    <outcome bullets…>                          (ignored)
    Speakers:                                    (or "Speaker:")
    Name, Role, Company
    Name, Role, Company
    Moderator: Name, Role, Company

Simple rows whose block is just a short label ("Welcome", "Coffee break",
"Lunch", "Cocktail", "Registration & Coffee") become section / break slides.

"To be confirmed" speakers/moderators are skipped (per project decision: no
placeholder slides until a real name lands).

Public API:
    parse_agenda_docx(path)  -> canonical dict
    parse_agenda_rows(rows)  -> canonical dict   (rows = list[(time, block_text)])
"""
from __future__ import annotations

import re
import unicodedata


# ── Session-type / break classification ──────────────────────────────────────

# Session-type words across EN / IT / ES / PL.
_TYPE_LINE_RE = re.compile(
    r"^\s*(discussion panel|panel|tavola rotonda|mesa redonda|panel de discusi[oó]n|"
    r"presentation|presentazione|presentaci[oó]n|prezentacja|"
    r"keynote|fireside chat|interview|intervista|entrevista|wywiad|"
    r"workshop|warsztaty|roundtable|round table|dibattito|debate|debata)\s*(\||$|\:|\sPanel)",
    re.IGNORECASE,
)
_LANG_RE = re.compile(r"(?:language|lingua|idioma|langue|j[ęe]zyk)\s*[:\-]\s*([A-Za-zÀ-ÿ]+)",
                      re.IGNORECASE)

# "Speakers:" / "Relatori:" / "Ponentes:" / … marker line.
_SPEAKERS_MARKER_RE = re.compile(
    r"^\s*(?:speakers?|relator[ei]|ponentes?|oradores?|intervengono|interviene|"
    r"prelegen[ct]i|referent[ein]?)\s*[:\-]?\s*$",
    re.IGNORECASE,
)
_MOD_PREFIX_RE = re.compile(
    r"^\s*(?:moderator|moderador|moderadora|moderador/a|moderatore|moderatrice|"
    r"moderator/ka|moderacja|modera)\s*[:\-]\s*(.*)$",
    re.IGNORECASE,
)

# Rows that are pure agenda furniture (no speakers) → section/break slides.
# Order matters: the first keyword found in the title wins, so put the more
# specific kinds (registration, welcome) before the generic "coffee"/"break".
_BREAK_KEYWORDS = {
    "registration": "registration", "registrazione": "registration",
    "accreditazioni": "registration", "accrediti": "registration",
    "registro": "registration", "registracion": "registration",
    "registración": "registration", "rejestracja": "registration",
    "welcome": "welcome", "apertura": "welcome", "benvenuto": "welcome",
    "bienvenida": "welcome", "powitanie": "welcome", "otwarcie": "welcome",
    "lunch": "lunch", "pranzo": "lunch", "almuerzo": "lunch", "comida": "lunch",
    "obiad": "lunch",
    "cocktail": "cocktail", "aperitivo": "cocktail", "cóctel": "cocktail",
    "coctel": "cocktail",
    "networking": "networking",
    "closing": "closing", "end of day": "closing", "fin del día": "closing",
    "clausura": "closing", "chiusura": "closing", "zakończenie": "closing",
    "coffee": "break", "caffè": "break", "caffe": "break", "café": "break",
    "cafe": "break", "pausa": "break", "przerwa": "break", "break": "break",
}

_TBC_RE = re.compile(
    r"^\s*to\s+be\s+confirmed\s*$|^\s*tbc\s*$|^\s*por\s+confirmar\s*$|"
    r"^\s*da\s+confermare\s*$|^\s*in\s+attesa\s+di\s+conferma\s*$",
    re.IGNORECASE)

# A bullet / outcome line usually starts with these or is a full sentence.
_BULLET_PREFIXES = ("•", "-", "·", "*", "–", "—", "▪", "●")

# Particles allowed lowercase inside a proper name.
_NAME_PARTICLES = {"de", "del", "van", "von", "le", "la", "el", "al", "y", "und",
                   "di", "da", "dos", "das", "bin", "binti", "mac", "mc", "ter", "den"}


def _strip(s: str) -> str:
    return (s or "").replace(" ", " ").strip()


def _is_tbc(s: str) -> bool:
    return bool(_TBC_RE.match(_strip(s)))


def _looks_like_name(s: str) -> bool:
    """Plausibly a person's name: 1–6 words, no digits, capitalised words."""
    s = _strip(s)
    if not s or len(s) > 70 or any(c.isdigit() for c in s):
        return False
    words = s.split()
    if not (1 <= len(words) <= 6):
        return False
    capitalised = 0
    for w in words:
        clean = w.strip("'\".()")
        if not clean:
            continue
        if clean.lower() in _NAME_PARTICLES:
            continue
        if clean[0].isupper() or not clean[0].isalpha():
            capitalised += 1
        else:
            return False
    return capitalised >= 1


# Moderator prefix in any language, incl. slash forms (Moderatore/trice,
# Moderador/a, Moderator/ka) the agenda generator emits.
_MOD_INLINE_RE = re.compile(
    r"^\s*moder[a-zà-ÿ]*(?:\s*/\s*[a-zà-ÿ]+)?\s*[:.\-]\s+(.+)$", re.IGNORECASE)


def _parse_speaker_line(line: str) -> dict | None:
    """'Name, Role, Company' OR 'Name — Role, Company' → dict, or None. Strips a
    leading 'Moderator:' (any language) and flags the speaker as moderator.

    Two agenda dialects: role separated from the name by a comma, or by a spaced
    em/en-dash / hyphen. The dash split is used only when the part before the dash
    has NO comma, so a mixed 'Name, Role — …, Company' line keeps just the name."""
    line = _strip(line)
    if not line or _is_tbc(line):
        return None

    is_mod = False
    mm = _MOD_INLINE_RE.match(line)
    while mm:                                  # repeat: some agendas double the prefix
        is_mod = True
        line = mm.group(1).strip()
        if not line or _is_tbc(line):
            return None
        mm = _MOD_INLINE_RE.match(line)

    md = re.match(r"^(.+?)\s+[—–-]\s+(.+)$", line)
    if md and "," not in md.group(1):
        name = md.group(1).strip()
        rest = md.group(2).strip()
        if _looks_like_name(name) and not _is_tbc(name):
            rp = [p.strip() for p in rest.split(",") if p.strip()]
            if len(rp) >= 2:
                return {"name": name, "role": ", ".join(rp[:-1]),
                        "company": rp[-1], "is_moderator": is_mod}
            if len(rp) == 1:
                return {"name": name, "role": "", "company": rp[0],
                        "is_moderator": is_mod}

    parts = [p.strip() for p in line.split(",")]
    name = parts[0]
    if _is_tbc(name) or not _looks_like_name(name):
        return None
    if len(parts) == 1:
        return {"name": name, "role": "", "company": "", "is_moderator": is_mod}
    company = parts[-1]
    role = ", ".join(parts[1:-1]) if len(parts) > 2 else ""
    return {"name": name, "role": role, "company": company, "is_moderator": is_mod}


# ── Time helpers ──────────────────────────────────────────────────────────────

_TIME_RE = re.compile(r"(\d{1,2})[:.h]?(\d{2})?")


def _norm_time(t: str) -> str:
    t = _strip(t).replace(" ", "")
    return t


def _start_time(t: str) -> str:
    """First HH:MM in a range like '9:15-10:05' → '9:15'."""
    t = _norm_time(t)
    first = re.split(r"[-–—]", t)[0]
    m = _TIME_RE.search(first)
    if not m:
        return first
    hh = m.group(1)
    mm = m.group(2) or "00"
    return f"{int(hh)}:{mm}"


def _slug(s: str) -> str:
    s = unicodedata.normalize("NFKD", s or "").encode("ascii", "ignore").decode("ascii")
    s = re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_").lower()
    return s or "x"


# ── Block classification ──────────────────────────────────────────────────────

def _classify_block(time: str, block: str, idx: int, day: int) -> dict | None:
    """Turn one table row (time + content block) into a canonical session dict."""
    raw_lines = [_strip(l) for l in block.replace("\r", "\n").split("\n")]
    lines = [l for l in raw_lines if l]
    if not lines:
        return None

    title = lines[0]

    # ── Pure section / break row (single short label, no speaker structure) ───
    low_title = title.lower()
    low_block = block.lower()
    _has_speaker_marker = any(mk in low_block for mk in
                              ("speaker", "relator", "ponente", "orador", "prelegen"))
    for kw, kind in _BREAK_KEYWORDS.items():
        if kw in low_title and len(lines) <= 2 and not _has_speaker_marker:
            return {
                "id": f"d{day}_{_slug(_start_time(time))}_{idx}",
                "day": day, "time": _norm_time(time), "start": _start_time(time),
                "type": "break", "break_kind": kind,
                "title": title, "title_2": None,
                "delivery_language": None, "speakers": [],
            }

    # ── Session type + delivery language from the "Type | Language:" line ──────
    stype = "session"
    delivery_lang = None
    for l in lines[1:4]:
        mt = _TYPE_LINE_RE.match(l)
        if mt:
            t = mt.group(1).lower()
            panel_words = ("panel", "roundtable", "round table", "tavola rotonda",
                           "mesa redonda", "dibattito", "debate", "debata")
            if any(w in t for w in panel_words):
                stype = "panel"
            else:
                stype = "presentation"
        ml = _LANG_RE.search(l)
        if ml:
            delivery_lang = ml.group(1).capitalize()

    # ── Speakers: every line after the Speaker(s): marker is one speaker; a
    #    Moderator line (anywhere) is also a speaker and must NOT stop collection.
    speakers: list[dict] = []
    in_speakers = False
    for l in lines[1:]:
        if _SPEAKERS_MARKER_RE.match(l):
            in_speakers = True
            continue
        is_mod_line = bool(_MOD_INLINE_RE.match(l))
        if not in_speakers and not is_mod_line:
            continue
        if l.startswith(_BULLET_PREFIXES):
            continue
        sp = _parse_speaker_line(l)
        if sp:
            speakers.append(sp)
            in_speakers = True       # a moderator line confirms we're in the people block

    # If no explicit "Speakers:" marker but the block clearly names people
    # (e.g. a short Presentation with a single "Name, Role, Company"), try the
    # tail lines as speakers.
    if not speakers and stype in ("session", "presentation"):
        for l in lines[1:]:
            if l.startswith(_BULLET_PREFIXES) or _TYPE_LINE_RE.match(l) or _LANG_RE.search(l):
                continue
            sp = _parse_speaker_line(l)
            if sp:
                speakers.append(sp)

    # Refine type: a "session" with multiple non-moderator speakers is a panel;
    # one speaker is a presentation. No named speakers → still a session shell
    # (e.g. "Presentation / To be confirmed") which the generator will skip.
    non_mod = [s for s in speakers if not s["is_moderator"]]
    if stype == "session":
        stype = "panel" if len(non_mod) >= 2 else "presentation"

    return {
        "id": f"d{day}_{_slug(_start_time(time))}_{idx}",
        "day": day, "time": _norm_time(time), "start": _start_time(time),
        "type": stype, "break_kind": None,
        "title": title, "title_2": None,
        "delivery_language": delivery_lang,
        "speakers": speakers,
    }


# ── Public entry points ───────────────────────────────────────────────────────

_DAY_MARKER_RE = re.compile(r"^\s*(?:day|d[ií]a|giorno|dzie[ńn])\s*(\d+)\s*$", re.IGNORECASE)


def parse_agenda_rows(rows: list[tuple[str, str]],
                      event_title: str = "") -> dict:
    """
    rows = list of (time_text, content_block_text) in document order.
    Returns the canonical agenda dict.
    """
    sessions: list[dict] = []
    day = 1
    for idx, (time, block) in enumerate(rows):
        # A row that is only a "Day N" marker bumps the day counter.
        dm = _DAY_MARKER_RE.match(_strip(block)) or _DAY_MARKER_RE.match(_strip(time))
        if dm:
            day = int(dm.group(1))
            continue
        if not _strip(block):
            continue
        sess = _classify_block(time, block, idx, day)
        if sess:
            sessions.append(sess)

    return {
        "event": {
            "title": _strip(event_title),
            "language": "en",
            "second_language": None,
        },
        "sessions": sessions,
    }


_URL_RE = re.compile(r"https?://\S+", re.IGNORECASE)


def _extract_register_url(doc) -> str | None:
    """Find a registration URL in the footers/paragraphs (RENMAD agendas put
    'Choose your pass here: https://…/register/' in the footer)."""
    texts = [p.text for p in doc.paragraphs]
    for sec in doc.sections:
        for hf in (sec.footer, sec.first_page_footer, sec.header, sec.first_page_header):
            texts += [p.text for p in hf.paragraphs]
    for t in texts:
        for m in _URL_RE.finditer(t or ""):
            url = m.group(0).rstrip(").,")
            if "register" in url.lower() or "pass" in (t or "").lower():
                return url
    for t in texts:
        m = _URL_RE.search(t or "")
        if m:
            return m.group(0).rstrip(").,")
    return None


def _extract_logo_wall(path: str):
    """Return the widest embedded image (the sponsor/speaker logo strip) as a PIL
    image, or None. The logo wall is much wider than tall (aspect > 2.5)."""
    import io
    import zipfile
    from PIL import Image
    best, best_aspect = None, 2.5
    try:
        z = zipfile.ZipFile(path)
    except Exception:
        return None
    for n in z.namelist():
        if "/media/" not in n:
            continue
        try:
            img = Image.open(io.BytesIO(z.read(n))).convert("RGBA")
        except Exception:
            continue
        aspect = img.width / max(1, img.height)
        if aspect > best_aspect and img.width >= 600:
            best, best_aspect = img, aspect
    return best


# Exact column-header labels emitted by the RENMAD Timed Agenda generator across
# its languages (en/es/it/pl/fr/de). Matched EXACTLY (not startswith) so a real
# session titled "Session 1 …" is never mistaken for a header row.
_HEADER_TIME_WORDS = {"time", "hora", "ora", "godzina", "heure", "zeit",
                      "horario", "orario", "tempo"}
_HEADER_PROG_WORDS = {"programme", "program", "programa", "programma", "programm"}
_HEADER_SPK_WORDS = {"speakers", "speaker", "relatori", "relatore", "ponentes",
                     "ponente", "prelegenci", "intervenants", "referenten",
                     "oradores", "relatori/trici"}

_TIME_CELL_RE = re.compile(r"\d{1,2}[:.h]\d{2}")


def _is_header_row(time: str, block: str) -> bool:
    """True for a table's column-header row (Time | Programme | Speakers)."""
    t = (time or "").strip().lower()
    b = (block or "").strip().lower().split("\n")[0].strip()
    return t in _HEADER_TIME_WORDS or b in _HEADER_PROG_WORDS


def _is_agenda_table(tbl) -> bool:
    """Identify THE agenda table(s) among all tables in the doc. A doc may also
    contain a logo+title header table and a footer/contact band table; with a
    header logo the agenda is no longer tables[0]. Multi-day / parallel-track
    agendas emit one agenda table per day per track. Detect by the column-header
    row, or (fallback) by most first-column cells looking like a time."""
    if not tbl.rows:
        return False
    c0 = tbl.rows[0].cells
    if len(c0) >= 2 and _is_header_row(_strip(c0[0].text), c0[1].text):
        return True
    timed = total = 0
    for row in tbl.rows:
        cells = row.cells
        if len(cells) < 2:
            continue
        total += 1
        if _TIME_CELL_RE.search(cells[0].text or ""):
            timed += 1
    return total >= 2 and timed >= max(2, total // 2)


def _first_title_text(doc, agenda_tables) -> str:
    """First substantial, non-label text in document order, stopping when the
    agenda table is reached. Handles both the header-logo (title in a table) and
    no-logo (title in the first paragraph) exports."""
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def _ok(t: str) -> bool:
        t = (t or "").strip()
        low = t.lower()
        return bool(t) and len(t) > 3 and "@" not in t \
            and "contact" not in low and "question" not in low \
            and not low.startswith("choose your pass")

    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            t = _strip(Paragraph(child, doc).text)
            if _ok(t):
                return t.split("\n")[0].strip()
        elif child.tag == qn("w:tbl"):
            if any(child is t._tbl for t in agenda_tables):
                break   # reached the agenda — the title (if any) came earlier
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    t = _strip(cell.text)
                    if _ok(t):
                        return t.split("\n")[0].strip()
    return ""


def _row_fields(row) -> "tuple | None":
    """Normalise a table row to (time, programme, speakers).

    Hand-built agendas use MERGED cells, which python-docx expands into repeated
    and empty cells (e.g. [time, time, '', '', prog, prog, speakers]). Collapse
    adjacent duplicates and drop empty fields, then map: first field = time,
    last field = speakers (when ≥3 fields remain), the richest middle field =
    programme. The clean 2-/3-column generator output passes through unchanged."""
    seq = []
    for cell in row.cells:
        s = _strip(cell.text)
        if seq and seq[-1][0] == s:        # adjacent duplicate from a merged cell
            continue
        seq.append((s, cell.text))
    seq = [(s, raw) for (s, raw) in seq if s]      # drop empty fields
    if not seq:
        return None
    time = seq[0][0]
    if len(seq) == 1:
        return (time, "", "")
    if len(seq) == 2:
        return (time, seq[1][1], "")
    speakers = seq[-1][1]
    programme = max(seq[1:-1], key=lambda x: len(x[0]))[1]
    return (time, programme, speakers)


def parse_agenda_docx(path: str) -> dict:
    """Open a .docx agenda (the RENMAD Timed Agenda generator's output) and parse
    it. Finds the agenda table(s) by signature — NOT by position — so it works
    whether or not a header logo pushes them past tables[0], and concatenates
    multiple agenda tables (multi-day / parallel tracks). Also extracts event
    meta (title, registration URL) + the sponsor logo wall."""
    import docx  # python-docx

    doc = docx.Document(path)

    agenda_tables = [t for t in doc.tables if _is_agenda_table(t)]

    # Event title = the first substantial text in DOCUMENT ORDER, before the
    # agenda starts. Walking in order (not paragraphs-then-tables) is essential:
    # with a header logo the title lives in a header TABLE cell and the first
    # paragraph is actually a day header ("Monday · 1 July"); without a logo the
    # title is the first paragraph. Either way the title comes first.
    event_title = _first_title_text(doc, agenda_tables)

    rows: list[tuple[str, str]] = []
    for tbl in agenda_tables:
        for row in tbl.rows:
            fields = _row_fields(row)
            if not fields:
                continue
            time, programme, speakers = fields
            # Skip the column-header row of each agenda table.
            if _is_header_row(time, programme):
                continue
            # The speaker column (rightmost) is folded in with a "Speakers:" marker
            # so the normal speaker parsing picks up each line as one speaker.
            block = programme
            if speakers:
                block = block.rstrip() + "\nSpeakers:\n" + speakers.strip()
            rows.append((time, block))

    ag = parse_agenda_rows(rows, event_title=event_title)
    ag["event"]["register_url"] = _extract_register_url(doc)
    ag["event"]["date_str"] = None      # not reliably in the docx body (header text box)
    ag["event"]["venue"] = None
    ag["logo_wall"] = _extract_logo_wall(path)
    return ag
